# import os
# import io
# import streamlit as st
# from dotenv import load_dotenv
# import google.generativeai as genai
# from PIL import Image

# # Load environment variables from .env file
# load_dotenv()

# # Configure the Gemini API client with your API key
# API_KEY = os.getenv("GOOGLE_API_KEY")
# if not API_KEY:
#     st.error("GOOGLE_API_KEY not found. Please set it in your .env file.")
# else:
#     genai.configure(api_key=API_KEY)

# # Initialize the Streamlit app interface
# st.set_page_config(page_title="Gemini Text Extractor", layout="centered")
# st.title("📄 Image Text Extractor")
# st.write("Upload one or more images, and the Gemini model will extract the text from each one.")
# st.markdown("---")

# # File uploader widget
# uploaded_files = st.file_uploader(
#     "Choose one or more image files...",
#     type=["jpg", "jpeg", "png"],
#     accept_multiple_files=True
# )

# if uploaded_files:
#     # Process files when the "Extract Text" button is clicked
#     if st.button("Extract Text"):
#         if not API_KEY:
#             st.warning("Please set your GOOGLE_API_KEY in the .env file to proceed.")
#             st.stop()
            
#         st.info("Extracting text from images...")
        
#         # Initialize the Gemini Vision model
#         try:
#             vision_model = genai.GenerativeModel('gemini-2.5-flash')
#         except Exception as e:
#             st.error(f"Failed to initialize Gemini model: {e}")
#             st.stop()

#         # Iterate through each uploaded file
#         for file in uploaded_files:
#             st.markdown(f"### Extracted from: `{file.name}`")
            
#             # Use a spinner while processing
#             with st.spinner('Working on it...'):
#                 try:
#                     # Read the image file into a bytes buffer
#                     image_bytes = file.read()
#                     image_stream = io.BytesIO(image_bytes)

#                     # Open the image using PIL (Pillow)
#                     image = Image.open(image_stream)

#                     # Define the prompt for the model
#                     prompt = "Extract all text from this image and provide it as a single block of text."

#                     # Send the image and prompt to the Gemini model
#                     response = vision_model.generate_content([prompt, image])

#                     # Extract the text and split it line by line
#                     extracted_text = response.text
#                     lines = extracted_text.split('\n')
                    
#                     # Display each line of text
#                     for line in lines:
#                         st.write(line)
                    
#                     st.markdown("---") # Separator for the next image

#                 except Exception as e:
#                     st.error(f"An error occurred while processing {file.name}: {e}")
#                     st.markdown("---")

#         st.success("Extraction complete!")



import os
import io
import re
import streamlit as st
from dotenv import load_dotenv
import google.generativeai as genai
from PIL import Image
from docx import Document  # For .docx export
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load environment variables
load_dotenv()

# Configure Gemini API
API_KEY = os.getenv("GOOGLE_API_KEY")
if not API_KEY:
    st.error("GOOGLE_API_KEY not found. Please set it in your .env file.")
else:
    genai.configure(api_key=API_KEY)

# Initialize Streamlit UI
st.set_page_config(page_title="📄 Organized Text Extractor", layout="centered")
st.title("📄 Image Text Extractor & Organizer")
st.write("Upload one or more image files. They will be processed in numerical order and exported to a Word document.")
st.markdown("---")

# File uploader
uploaded_files = st.file_uploader(
    "Choose one or more image files...",
    type=["jpg", "jpeg", "png"],
    accept_multiple_files=True
)

# Global variable to store extracted text
extracted_data = []

if uploaded_files:
    # Sort files by number in filename: 1.jpg, 2.jpg, etc.
    def extract_number(file):
        match = re.search(r'(\d+)', file.name)
        return int(match.group(1)) if match else float('inf')  # Unnumbered goes last

    sorted_files = sorted(uploaded_files, key=extract_number)

    st.write("📁 Files will be processed in this order:")
    for i, file in enumerate(sorted_files):
        num_match = re.search(r'(\d+)', file.name)
        q_num = f"Q{num_match.group(1)}" if num_match else "Unknown"
        st.write(f"{i+1}. `{file.name}` → {q_num}")

    st.markdown("---")

    # Extract Text Button
    if st.button("🔍 Extract Text in Order"):
        if not API_KEY:
            st.warning("Please set GOOGLE_API_KEY in .env")
            st.stop()

        try:
            model = genai.GenerativeModel('gemini-2.5-flash')  # Correct model name
        except Exception as e:
            st.error(f"Failed to initialize Gemini model: {e}")
            st.stop()

        st.info("Extracting text from images in order...")

        # Clear previous data
        extracted_data.clear()

        # Process each file in sorted order
        for file in sorted_files:
            with st.spinner(f"Processing {file.name}..."):
                try:
                    # Read image
                    image_bytes = file.read()
                    image_stream = io.BytesIO(image_bytes)
                    image = Image.open(image_stream)

                    # Prompt
                    prompt = "Extract all visible text from this image exactly as it appears. Preserve line breaks, bilingual content (Nepali + English), and formatting."

                    # Generate content
                    response = model.generate_content([prompt, image])
                    full_text = response.text.strip() if response.text else "[No text found]"

                    # Store result
                    extracted_data.append({
                        "filename": file.name,
                        "text": full_text
                    })

                    # Display immediately
                    st.markdown(f"### 📄 Extracted from: `{file.name}`")
                    st.write(full_text)
                    st.markdown("<hr>", unsafe_allow_html=True)

                except Exception as e:
                    error_msg = f"Error processing {file.name}: {str(e)}"
                    st.error(error_msg)
                    extracted_data.append({
                        "filename": file.name,
                        "text": f"[Error: {e}]"
                    })
                    st.markdown("<hr>", unsafe_allow_html=True)

        st.success("✅ Extraction complete! Ready to download.")

        # --- Generate .docx File ---
        doc = Document()
        doc.add_heading('Extracted Exam Paper', 0)

        for item in extracted_data:
            # Add filename as heading
            p = doc.add_paragraph()
            runner = p.add_run(f"From: {item['filename']}")
            runner.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Add extracted text
            text_lines = item["text"].split("\n")
            for line in text_lines:
                doc.add_paragraph(line.strip())

            # Add spacing
            doc.add_paragraph("_" * 50)  # Separator line

        # Save to BytesIO buffer
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)

        # Download Button
        st.download_button(
            label="⬇️ Download as Word (.docx)",
            data=bio,
            file_name="extracted_exam_paper.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
